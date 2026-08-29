import os
import sys
import numpy as np
import matplotlib.pyplot as plt
import docx
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

# ---------------------------------------------------------
# STYLING HELPERS
# ---------------------------------------------------------
def set_cell_margins(cell, top=100, bottom=100, left=150, right=150):
    tcPr = cell._element.get_or_add_tcPr()
    tcMar = OxmlElement('w:tcMar')
    for m, val in [('w:top', top), ('w:bottom', bottom), ('w:left', left), ('w:right', right)]:
        node = OxmlElement(m)
        node.set(qn('w:w'), str(val))
        node.set(qn('w:type'), 'dxa')
        tcMar.append(node)
    tcPr.append(tcMar)

def set_cell_shading(cell, color_hex):
    tcPr = cell._element.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), color_hex)
    tcPr.append(shd)

# ---------------------------------------------------------
# MAIN ANALYSIS AND GENERATION
# ---------------------------------------------------------
def run_combined_analysis():
    print("🚀 Initializing Combined Model Performance Analyzer...")
    os.makedirs("combined_outputs", exist_ok=True)
    
    # 1. DEFINE SOURCE DATA (Metrics collected across all previous runs)
    metrics = {
        "Chatbot (Greeting)": {
            "Test 1 (Ideal)": [1.0000, 1.0000, 1.0000, 1.0000],
            "Test 2 (Stress)": [0.8750, 1.0000, 0.6250, 0.7692],
            "Test 3 (Production)": [0.9167, 0.8000, 1.0000, 0.8889],
            "Golden Test": [1.0000, 1.0000, 1.0000, 1.0000]
        },
        "Chatbot (Off-Topic)": {
            "Test 1 (Ideal)": [1.0000, 1.0000, 1.0000, 1.0000],
            "Test 2 (Stress)": [0.5417, 0.2000, 0.1250, 0.1538],
            "Test 3 (Production)": [1.0000, 1.0000, 1.0000, 1.0000],
            "Golden Test": [1.0000, 1.0000, 1.0000, 1.0000]
        },
        "Roommate Matcher": {
            "Test 1 (Ideal)": [1.0000, 1.0000, 1.0000, 1.0000],
            "Test 2 (Stress)": [0.8000, 0.7143, 1.0000, 0.8333],
            "Test 3 (Production)": [0.8667, 0.8000, 1.0000, 0.8889],
            "Golden Test": [0.7500, 0.6667, 1.0000, 0.8000]
        },
        "Auto-Tagging Model": {
            "Test 1 (Ideal)": [0.9732, 0.9412, 0.8889, 0.9143],
            "Test 2 (Stress)": [0.5804, 0.2333, 0.9333, 0.3733],
            "Test 3 (Production)": [0.6000, 0.2432, 1.0000, 0.3913],
            "Golden Test": [0.5893, 0.2500, 0.9375, 0.3947]
        }
    }
    
    # 2. COMPUTE MATHEMATICAL AVERAGES
    averages = {}
    for model, runs in metrics.items():
        data_matrix = np.array(list(runs.values())) # shape (4, 4)
        avg_scores = np.mean(data_matrix, axis=0) # [Accuracy, Precision, Recall, F1]
        averages[model] = avg_scores.tolist()
        
    print("📈 Calculated Model Averages successfully.")
    
    # 3. PLOT PERFORMANCE COMPARISON CHART
    print("📊 Generating F1-Score comparison bar chart...")
    models_list = list(metrics.keys())
    runs_labels = ["Test 1 (Ideal)", "Test 2 (Stress)", "Test 3 (Prod)", "Golden Test", "Average"]
    
    # Extract F1 scores (index 3)
    f1_data = []
    for model in models_list:
        model_f1s = [metrics[model][r][3] * 100 for r in ["Test 1 (Ideal)", "Test 2 (Stress)", "Test 3 (Production)", "Golden Test"]]
        model_f1s.append(averages[model][3] * 100) # Append average F1
        f1_data.append(model_f1s)
        
    f1_data = np.array(f1_data) # shape (4, 5)
    
    fig, ax = plt.subplots(figsize=(10, 6))
    x = np.arange(len(models_list))
    width = 0.15
    
    colors = ['#93C5FD', '#F87171', '#60A5FA', '#34D399', '#10B981'] # Slate colors
    
    for idx, label in enumerate(runs_labels):
        ax.bar(x + idx * width - (len(runs_labels)*width)/2 + width/2, f1_data[:, idx], width, label=label, color=colors[idx])
        
    ax.set_ylabel('F1 Score (%)')
    ax.set_title('Nestify AI Models: F1 Score Comparison Across Evaluation Runs')
    ax.set_xticks(x)
    ax.set_xticklabels(models_list)
    ax.set_ylim(0, 110)
    ax.legend(loc='lower left')
    ax.grid(axis='y', linestyle='--', alpha=0.5)
    
    plt.tight_layout()
    chart_path = "combined_outputs/f1_comparison.png"
    plt.savefig(chart_path, dpi=150)
    plt.close()
    
    # 4. WRITE THE WORD DOCUMENT
    print("📝 Compiling Combined Evaluation Report into DOCX...")
    doc = docx.Document()
    
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)
        
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Arial'
    font.size = Pt(10.5)
    
    # Document Cover Title
    title_p = doc.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title_p.add_run("Nestify AI Microservices\nCombined Performance Analysis & Comparative Report")
    title_run.font.name = 'Arial'
    title_run.font.size = Pt(20)
    title_run.font.bold = True
    title_run.font.color.rgb = RGBColor(30, 41, 59) # Slate 800
    
    meta_p = doc.add_paragraph()
    meta_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    meta_run = meta_p.add_run("Multi-Scenario Benchmark Analysis | Prepared by: Nestify AI Team")
    meta_run.font.size = Pt(9.5)
    meta_run.font.italic = True
    meta_run.font.color.rgb = RGBColor(100, 116, 139)
    
    doc.add_paragraph("").paragraph_format.space_before = Pt(10)
    
    # Executive Summary
    h1 = doc.add_paragraph()
    h1.paragraph_format.space_before = Pt(12)
    h1.paragraph_format.space_after = Pt(6)
    r = h1.add_run("Executive Summary")
    r.font.size = Pt(14)
    r.font.bold = True
    r.font.color.rgb = RGBColor(37, 99, 235) # Blue 600
    
    summary_text = (
        "This report delivers a unified performance analysis of the Nestify platform's AI microservices. "
        "By synthesizing results across four distinct evaluation paradigms—ranging from idealized code scenarios to stress-tested API outages—we "
        "compute the mathematically grounded performance average for each service. This combined benchmark serves as our primary baseline "
        "for tracking production reliability and predicting user-experience quality under diverse operating environments."
    )
    doc.add_paragraph(summary_text)
    
    # ---------------------------------------------------------
    # COMPARATIVE PERFORMANCE TABLE
    # ---------------------------------------------------------
    h1 = doc.add_paragraph()
    h1.paragraph_format.space_before = Pt(16)
    h1.paragraph_format.space_after = Pt(6)
    r = h1.add_run("1. Performance Metrics Across Evaluation Runs")
    r.font.size = Pt(14)
    r.font.bold = True
    r.font.color.rgb = RGBColor(37, 99, 235)
    
    doc.add_paragraph("The table below details accuracy, precision, recall, and F1 scores across all four individual runs, including the mathematically calculated average.")
    
    table = doc.add_table(rows=1, cols=6)
    table.alignment = docx.enum.table.WD_TABLE_ALIGNMENT.CENTER
    
    hdr_cells = table.rows[0].cells
    headers = ["Service & Run", "Accuracy", "Precision", "Recall", "F1 Score", "Status"]
    for i, title_text in enumerate(headers):
        hdr_cells[i].text = title_text
        set_cell_margins(hdr_cells[i])
        set_cell_shading(hdr_cells[i], "1E3A8A") # Navy Blue
        for run in hdr_cells[i].paragraphs[0].runs:
            run.font.bold = True
            run.font.color.rgb = RGBColor(255, 255, 255)
            
    # Populate data row by row
    row_idx = 0
    for model in models_list:
        # Test 1
        row = table.add_row().cells
        row[0].text = f"{model} - Test 1"
        row[1].text = f"{metrics[model]['Test 1 (Ideal)'][0]*100:.2f}%"
        row[2].text = f"{metrics[model]['Test 1 (Ideal)'][1]*100:.2f}%"
        row[3].text = f"{metrics[model]['Test 1 (Ideal)'][2]*100:.2f}%"
        row[4].text = f"{metrics[model]['Test 1 (Ideal)'][3]*100:.2f}%"
        row[5].text = "Ideal"
        set_cell_shading(row[0], "F8FAFC")
        
        # Test 2
        row = table.add_row().cells
        row[0].text = f"{model} - Test 2"
        row[1].text = f"{metrics[model]['Test 2 (Stress)'][0]*100:.2f}%"
        row[2].text = f"{metrics[model]['Test 2 (Stress)'][1]*100:.2f}%"
        row[3].text = f"{metrics[model]['Test 2 (Stress)'][2]*100:.2f}%"
        row[4].text = f"{metrics[model]['Test 2 (Stress)'][3]*100:.2f}%"
        row[5].text = "Rate Limit"
        set_cell_shading(row[0], "FEF2F2") # Red tint
        
        # Test 3
        row = table.add_row().cells
        row[0].text = f"{model} - Test 3"
        row[1].text = f"{metrics[model]['Test 3 (Production)'][0]*100:.2f}%"
        row[2].text = f"{metrics[model]['Test 3 (Production)'][1]*100:.2f}%"
        row[3].text = f"{metrics[model]['Test 3 (Production)'][2]*100:.2f}%"
        row[4].text = f"{metrics[model]['Test 3 (Production)'][3]*100:.2f}%"
        row[5].text = "Production"
        set_cell_shading(row[0], "EFF6FF") # Blue tint
        
        # Golden
        row = table.add_row().cells
        row[0].text = f"{model} - Golden"
        row[1].text = f"{metrics[model]['Golden Test'][0]*100:.2f}%"
        row[2].text = f"{metrics[model]['Golden Test'][1]*100:.2f}%"
        row[3].text = f"{metrics[model]['Golden Test'][2]*100:.2f}%"
        row[4].text = f"{metrics[model]['Golden Test'][3]*100:.2f}%"
        row[5].text = "Gold Standard"
        set_cell_shading(row[0], "ECFDF5") # Green tint
        
        # Average (BOLD)
        row = table.add_row().cells
        row[0].text = f"{model} - AVERAGE"
        row[1].text = f"{averages[model][0]*100:.2f}%"
        row[2].text = f"{averages[model][1]*100:.2f}%"
        row[3].text = f"{averages[model][2]*100:.2f}%"
        row[4].text = f"{averages[model][3]*100:.2f}%"
        row[5].text = "Computed Mean"
        set_cell_shading(row[0], "F1F5F9")
        for cell in row:
            for para in cell.paragraphs:
                for run in para.runs:
                    run.font.bold = True
                    
        # Apply padding to all cell blocks
        for row_cells in table.rows:
            for cell in row_cells.cells:
                set_cell_margins(cell)
                
    # ---------------------------------------------------------
    # CHART EMBEDDING
    # ---------------------------------------------------------
    doc.add_page_break()
    h1 = doc.add_paragraph()
    h1.paragraph_format.space_before = Pt(12)
    h1.paragraph_format.space_after = Pt(6)
    r = h1.add_run("2. Graphical Comparison (F1-Score Profile)")
    r.font.size = Pt(14)
    r.font.bold = True
    r.font.color.rgb = RGBColor(37, 99, 235)
    
    p_img = doc.add_paragraph()
    p_img.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_img.add_run().add_picture(chart_path, width=Inches(6))
    caption = doc.add_paragraph("Figure 2.1: Bar Chart comparing F1 Scores of all four test scenarios against the final average.")
    caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
    caption.runs[0].font.size = Pt(9)
    caption.runs[0].font.italic = True
    
    # ---------------------------------------------------------
    # IN-DEPTH COMPARATIVE EXPLANATION
    # ---------------------------------------------------------
    h1 = doc.add_paragraph()
    h1.paragraph_format.space_before = Pt(16)
    h1.paragraph_format.space_after = Pt(6)
    r = h1.add_run("3. In-Depth Comparative Analysis")
    r.font.size = Pt(14)
    r.font.bold = True
    r.font.color.rgb = RGBColor(37, 99, 235)
    
    analysis_paragraphs = [
        "**Test 1 (Ideal Baseline)** used simple, direct query variations and highly polarized roommate profiles (either 100% identical or completely opposite genders). "
        "This led to a deceptive 100% score across almost all metrics. While it proved the core vector logic and intent routers worked, it did not reflect real production environments.",
        
        "**Test 2 (Stress Test)** challenged the models under critical operating conditions. By introducing extreme borderlines (e.g., roommate candidates sharing most attributes but possessing conflicting smoking habits) "
        "and evaluating chatbot queries while triggering Gemini free-tier rate limits, we observed natural degradations. The chatbot's F1 score fell to 15.38% as it hit rate-limit fallbacks. "
        "This demonstrated the system's resilience but also highlighted the impact of API availability.",
        
        "**Test 3 (Production Benchmark)** represented typical student behavior under stable API conditions. Bypassing quota limits allowed the chatbot's intent classification F1 score to return to 100%. "
        "By setting the roommate matcher threshold to a realistic 65% similarity boundary, we captured acceptable preference variations, resulting in a healthy 86.67% accuracy.",
        
        "**Golden Dataset Test** validated the system against human-curated ground truths. This run confirmed high conversational precision (100%) but highlighted that vector roommate matchmaking requires "
        "ongoing preference-vector normalization, as accuracy settled at 75.00%.",
        
        "**Computed Average** represents the expected mean performance across various operating realities (stable API, rate-limiting, noise, clear-cut cases). "
        "For example, the Chatbot (Greeting) maintains a solid F1 Average of 91.45%, validating its high readiness for production."
    ]
    
    for p_text in analysis_paragraphs:
        doc.add_paragraph(p_text)
        
    doc_path = "combined_evaluation_report.docx"
    doc.save(doc_path)
    print(f"✅ Combined analysis saved to {doc_path}")

if __name__ == "__main__":
    run_combined_analysis()
